import tkinter as tk
from tkinter import scrolledtext, filedialog
import requests
from scapy.all import ARP, Ether, srp
import socket
from multiprocessing import Pool
import nvdlib 
from requests.exceptions import ReadTimeout
import threading
from docx import Document

api_key = "01hy8tnhzb6k91zfmc48hesx3h01hy8tp95f6d006dts2pmjnyyslfzdnhovlnwp"

current_text_area = None  # Global variable to keep track of the current text area

def mac_lookup(mac_address):
    url = f"https://api.maclookup.app/v2/macs/{mac_address}"
    headers = {"X-Authentication-Token": api_key}
    try:
        response = requests.get(url, headers=headers, timeout=10)
        if response.status_code == 200:
            data = response.json()
            return data['company']
        else:
            return "unknown"
    except ReadTimeout:
        return "unknown (timeout)"

def scan_network(ip_range, text_area):
    devices = []
    arp = ARP(pdst=ip_range)   #Finds Ip and MAC with ARP request
    ether = Ether(dst="ff:ff:ff:ff:ff:ff")
    packet = ether/arp
    result = srp(packet, timeout=3, verbose=False)[0]
    for sent, received in result:
        devices.append({'ip': received.psrc, 'mac': received.hwsrc})
    text_area.delete('1.0', tk.END)  # Clear the text area
    text_area.insert(tk.END, "Devices on the network:\n")
    text_area.insert(tk.END, "IP Address\t\tMAC Address\t\tVendor\n")
    text_area.insert(tk.END, "-----------------------------------------\n")
    for device in devices:
        mac_vendor = mac_lookup(device['mac'])
        text_area.insert(tk.END, f"{device['ip']}\t\t{device['mac']}\t\t{mac_vendor}\n")

def start_scan_network(text_area):
    ip_range = '192.168.51.0/24'  # Replace with the appropriate IP range
    threading.Thread(target=scan_network, args=(ip_range, text_area)).start() # Starting in seprate Thread to stop hanging during scan 

def check_vulnerabilities(vendor_entry, text_area):
    vendor_name = vendor_entry.get()
    text_area.delete('1.0', tk.END)  # Clear the text area
    text_area.insert(tk.END, f"Checking for vulnerabilities for vendor: {vendor_name}...\n")
    threading.Thread(target=search_vulnerabilities, args=(vendor_name, text_area)).start()

def search_vulnerabilities(vendor_name, text_area):
    try:
        r = nvdlib.searchCVE_V2(keywordSearch=vendor_name, limit=100)
        for oneCVE in r:
            CV = oneCVE.descriptions
            result = CV[0].value
            text_area.insert(tk.END, result + "\n\n")
    except ReadTimeout:
        text_area.insert(tk.END, "Error: Request timed out.\n")

def start_port_scan(ip_entry, text_area):
    target_ip = ip_entry.get()
    text_area.delete('1.0', tk.END)  # Clear the text area
    text_area.insert(tk.END, f"Scanning ports on {target_ip}...\n")

    # Run the port scan in a separate thread
    threading.Thread(target=perform_port_scan, args=(target_ip, text_area)).start()

def perform_port_scan(target_ip, text_area):
    port_range = range(1, 1001)
    pool = Pool(processes=50)
    open_ports = pool.map(scan_port, [(port, target_ip) for port in port_range])
    open_ports = [port for port in open_ports if port is not None]
    text_area.insert(tk.END, f"Open ports on {target_ip}:\n")
    if open_ports:
        for port in open_ports:
            text_area.insert(tk.END, f"Port {port} is open\n")
    else:
        text_area.insert(tk.END, "No open ports found\n")

def scan_port(args):
    port, target_ip = args
    try:
        sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        sock.settimeout(0.5)
        result = sock.connect_ex((target_ip, port))
        if result == 0:
            return port
    except Exception as e:
        print(f"Error scanning port {port}: {e}")
    finally:
        sock.close()
    return None

def show_frame(frame, text_area):
    global current_text_area
    current_text_area = text_area
    frame.tkraise()

def save_report_to_file(text_area_scan, text_area_ports, text_area_vuln):
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word documents", "*.docx"), ("All files", "*.*")])
    if file_path:
        document = Document()
        document.add_heading('Network Scanning Report', 0)
        
        document.add_heading('Scanning Results', level=1)
        document.add_paragraph(text_area_scan.get("1.0", tk.END))
        
        document.add_heading('Port Scanning Results', level=1)
        document.add_paragraph(text_area_ports.get("1.0", tk.END))
        
        document.add_heading('Vulnerability Results', level=1)
        document.add_paragraph(text_area_vuln.get("1.0", tk.END))
        
        document.save(file_path)

def gui():
    root = tk.Tk()
    root.title("Network Scanner")
    root.geometry("1300x800")
    root.configure(bg='#f0f0f0')

    scanning_frame = tk.Frame(root, bg='#f0f0f0')
    ports_frame = tk.Frame(root, bg='#f0f0f0')
    vulnerabilities_frame = tk.Frame(root, bg='#f0f0f0')

    for frame in (scanning_frame, ports_frame, vulnerabilities_frame):
        frame.grid(row=1, column=0, sticky='news')

    button_frame = tk.Frame(root, bg='#333')
    button_frame.grid(row=0, column=0, pady=10)

    button_style = {
        'bg': '#4CAF50',
        'fg': 'white',
        'font': ('Arial', 12),
        'padx': 20,
        'pady': 10
    }

    scan_button = tk.Button(button_frame, text="Scanning", command=lambda: show_frame(scanning_frame, text_area_scan), **button_style)
    scan_button.pack(side='left', padx=10)

    ports_button = tk.Button(button_frame, text="Ports", command=lambda: show_frame(ports_frame, text_area_ports), **button_style)
    ports_button.pack(side='left', padx=10)

    vulnerabilities_button = tk.Button(button_frame, text="Vulnerabilities", command=lambda: show_frame(vulnerabilities_frame, text_area_vuln), **button_style)
    vulnerabilities_button.pack(side='left', padx=10)

    download_button = tk.Button(button_frame, text="Download Report", command=lambda: save_report_to_file(text_area_scan, text_area_ports, text_area_vuln), **button_style)
    download_button.pack(side='left', padx=10)

    # Scanning Frame
    text_area_scan = scrolledtext.ScrolledText(scanning_frame, width=150, height=20, bg='#e8e8e8', font=('Arial', 10))
    text_area_scan.pack(padx=10, pady=10)
    start_scan_button = tk.Button(scanning_frame, text="Start Scan", command=lambda: start_scan_network(text_area_scan), **button_style)
    start_scan_button.pack(pady=5)

    # Ports Frame
    ip_label = tk.Label(ports_frame, text="Enter IP Address:", bg='#f0f0f0', font=('Arial', 12))
    ip_label.pack(pady=5)
    ip_entry = tk.Entry(ports_frame, font=('Arial', 12))
    ip_entry.pack(pady=5)
    text_area_ports = scrolledtext.ScrolledText(ports_frame, width=150, height=20, bg='#e8e8e8', font=('Arial', 10))
    text_area_ports.pack(padx=10, pady=10)
    start_port_scan_button = tk.Button(ports_frame, text="Start Port Scan", command=lambda: start_port_scan(ip_entry, text_area_ports), **button_style)
    start_port_scan_button.pack(pady=5)

    # Vulnerabilities Frame
    vendor_label = tk.Label(vulnerabilities_frame, text="Enter Vendor Name:", bg='#f0f0f0', font=('Arial', 12))
    vendor_label.pack(pady=5)
    vendor_entry = tk.Entry(vulnerabilities_frame, font=('Arial', 12))
    vendor_entry.pack(pady=5)
    text_area_vuln = scrolledtext.ScrolledText(vulnerabilities_frame, width=150, height=20, bg='#e8e8e8', font=('Arial', 10))
    text_area_vuln.pack(padx=10, pady=10)
    check_vuln_button = tk.Button(vulnerabilities_frame, text="Check Vulnerabilities", command=lambda: check_vulnerabilities(vendor_entry, text_area_vuln), **button_style)
    check_vuln_button.pack(pady=5)

    show_frame(scanning_frame, text_area_scan)  # Show the scanning frame by default
    root.mainloop()

if __name__ == "__main__":
    gui()
